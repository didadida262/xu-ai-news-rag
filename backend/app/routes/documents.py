from flask import request, current_app
from flask_restx import Namespace, Resource, fields
from flask_jwt_extended import jwt_required
from datetime import datetime
from werkzeug.utils import secure_filename
import sys
import os
import json
import pandas as pd
from pathlib import Path

# 添加项目路径
backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
if backend_dir not in sys.path:
    sys.path.insert(0, backend_dir)

from models.document import Document
from models.database import db

documents_ns = Namespace('documents', description='文档管理相关操作')

# 数据模型
document_model = documents_ns.model('Document', {
    'id': fields.Integer(description='文档ID'),
    'title': fields.String(description='标题'),
    'content': fields.String(description='内容（摘要）'),
    'summary': fields.String(description='摘要'),
    'source_type': fields.String(description='来源类型: rss/web'),
    'source_url': fields.String(description='来源URL'),
    'source_name': fields.String(description='来源名称'),
    'tags': fields.List(fields.String, description='标签'),
    'metadata': fields.Raw(description='元数据'),
    'is_processed': fields.Boolean(description='是否已处理'),
    'is_vectorized': fields.Boolean(description='是否已向量化'),
    'created_at': fields.String(description='创建时间'),
    'updated_at': fields.String(description='更新时间')
})

@documents_ns.route('')
class DocumentsList(Resource):
    @jwt_required()
    @documents_ns.doc(params={
        'page': '页码（默认1）',
        'per_page': '每页数量（默认20）',
        'source_type': '来源类型筛选（rss/web）',
        'search': '搜索关键词（标题或内容）',
        'start_date': '开始日期（YYYY-MM-DD）',
        'end_date': '结束日期（YYYY-MM-DD）',
        'is_processed': '是否已处理（true/false）',
        'is_vectorized': '是否已向量化（true/false）'
    })
    def get(self):
        """获取文档列表（支持分页、搜索、筛选）"""
        try:
            # 获取查询参数
            page = request.args.get('page', 1, type=int)
            per_page = request.args.get('per_page', 20, type=int)
            source_type = request.args.get('source_type')
            search = request.args.get('search')
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            is_processed = request.args.get('is_processed')
            is_vectorized = request.args.get('is_vectorized')
            
            # 构建查询
            query = Document.query
            
            # 来源类型筛选
            if source_type:
                query = query.filter(Document.source_type == source_type)
            
            # 日期范围筛选
            if start_date:
                try:
                    start = datetime.strptime(start_date, '%Y-%m-%d')
                    query = query.filter(Document.created_at >= start)
                except ValueError:
                    pass
            
            if end_date:
                try:
                    end = datetime.strptime(end_date, '%Y-%m-%d')
                    # 设置为当天的结束时间
                    end = end.replace(hour=23, minute=59, second=59)
                    query = query.filter(Document.created_at <= end)
                except ValueError:
                    pass
            
            # 搜索
            if search:
                query = query.filter(
                    db.or_(
                        Document.title.contains(search),
                        Document.content.contains(search),
                        Document.summary.contains(search)
                    )
                )
            
            # 处理状态筛选
            if is_processed is not None:
                is_processed_bool = is_processed.lower() == 'true'
                query = query.filter(Document.is_processed == is_processed_bool)
            
            # 向量化状态筛选
            if is_vectorized is not None:
                is_vectorized_bool = is_vectorized.lower() == 'true'
                query = query.filter(Document.is_vectorized == is_vectorized_bool)
            
            # 分页
            pagination = query.order_by(Document.created_at.desc()).paginate(
                page=page, per_page=per_page, error_out=False
            )
            
            # 记录查询日志
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"文档列表查询: 总数={pagination.total}, 当前页={page}, 每页={per_page}, 筛选条件: source_type={source_type}, search={search}, is_processed={is_processed}, is_vectorized={is_vectorized}")
            
            return {
                'items': [doc.to_dict() for doc in pagination.items],
                'pagination': {
                    'page': page,
                    'per_page': per_page,
                    'total': pagination.total,
                    'pages': pagination.pages,
                    'has_prev': pagination.has_prev,
                    'has_next': pagination.has_next
                }
            }, 200
            
        except Exception as e:
            return {'error': f'获取文档列表失败: {str(e)}'}, 500


@documents_ns.route('/<int:doc_id>')
class DocumentDetail(Resource):
    @jwt_required()
    @documents_ns.marshal_with(document_model)
    def get(self, doc_id):
        """获取文档详情（包含完整内容）"""
        try:
            doc = Document.query.get_or_404(doc_id)
            return doc.to_dict_full(), 200
        except Exception as e:
            return {'error': f'获取文档详情失败: {str(e)}'}, 500
    
    @jwt_required()
    @documents_ns.expect(document_model)
    @documents_ns.marshal_with(document_model)
    def put(self, doc_id):
        """更新文档元数据（标签、来源等）"""
        try:
            doc = Document.query.get_or_404(doc_id)
            data = request.get_json()
            
            # 更新可编辑的字段
            if 'tags' in data:
                doc.tags = data['tags']
            if 'source_name' in data:
                doc.source_name = data['source_name']
            if 'source_url' in data:
                doc.source_url = data['source_url']
            if 'title' in data:
                doc.title = data['title']
            if 'summary' in data:
                doc.summary = data['summary']
            
            doc.updated_at = datetime.utcnow()
            db.session.commit()
            
            return doc.to_dict_full(), 200
        except Exception as e:
            db.session.rollback()
            return {'error': f'更新文档失败: {str(e)}'}, 500
    
    @jwt_required()
    def delete(self, doc_id):
        """删除文档"""
        try:
            doc = Document.query.get_or_404(doc_id)
            db.session.delete(doc)
            db.session.commit()
            return {'message': '文档删除成功'}, 200
        except Exception as e:
            db.session.rollback()
            return {'error': f'删除文档失败: {str(e)}'}, 500


@documents_ns.route('/batch-delete')
class BatchDeleteDocuments(Resource):
    @jwt_required()
    @documents_ns.doc(params={
        'ids': '文档ID列表，用逗号分隔（如：1,2,3）'
    })
    def post(self):
        """批量删除文档"""
        try:
            ids_str = request.args.get('ids') or request.get_json().get('ids', '')
            if not ids_str:
                return {'error': '请提供要删除的文档ID列表'}, 400
            
            # 解析ID列表
            try:
                ids = [int(id.strip()) for id in ids_str.split(',') if id.strip()]
            except ValueError:
                return {'error': '文档ID格式错误'}, 400
            
            if not ids:
                return {'error': '请提供有效的文档ID'}, 400
            
            # 查找并删除文档
            documents = Document.query.filter(Document.id.in_(ids)).all()
            deleted_count = len(documents)
            
            for doc in documents:
                db.session.delete(doc)
            
            db.session.commit()
            
            return {
                'message': f'成功删除 {deleted_count} 个文档',
                'deleted_count': deleted_count
            }, 200
            
        except Exception as e:
            db.session.rollback()
            return {'error': f'批量删除失败: {str(e)}'}, 500


@documents_ns.route('/stats')
class DocumentStats(Resource):
    @jwt_required()
    def get(self):
        """获取文档统计信息"""
        try:
            stats = Document.get_stats()
            return stats, 200
        except Exception as e:
            return {'error': f'获取统计信息失败: {str(e)}'}, 500


@documents_ns.route('/upload')
class UploadDocument(Resource):
    @jwt_required()
    def post(self):
        """上传文件到知识库（支持多种文件类型）"""
        try:
            if 'file' not in request.files:
                return {'error': '未提供文件'}, 400
            
            file = request.files['file']
            if file.filename == '':
                return {'error': '文件名为空'}, 400
            
            # 获取上传配置
            upload_folder = current_app.config.get('UPLOAD_FOLDER', './uploads')
            os.makedirs(upload_folder, exist_ok=True)
            
            # 保存文件
            filename = secure_filename(file.filename)
            file_path = os.path.join(upload_folder, filename)
            file.save(file_path)
            file_size = os.path.getsize(file_path)
            
            # 根据文件类型解析内容
            file_ext = Path(filename).suffix.lower()
            title = Path(filename).stem
            content = ''
            tags = []
            
            try:
                if file_ext in ['.txt', '.md']:
                    # 文本文件
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                elif file_ext in ['.xlsx', '.xls']:
                    # Excel文件
                    df = pd.read_excel(file_path)
                    # 将Excel转换为文本
                    content = df.to_string(index=False)
                    # 提取列名作为标签
                    tags = list(df.columns) if len(df.columns) > 0 else []
                
                elif file_ext == '.csv':
                    # CSV文件
                    df = pd.read_csv(file_path, encoding='utf-8')
                    content = df.to_string(index=False)
                    tags = list(df.columns) if len(df.columns) > 0 else []
                
                elif file_ext == '.html':
                    # HTML文件
                    from bs4 import BeautifulSoup
                    with open(file_path, 'r', encoding='utf-8') as f:
                        soup = BeautifulSoup(f.read(), 'html.parser')
                        # 提取文本内容
                        for script in soup(['script', 'style']):
                            script.decompose()
                        content = soup.get_text(separator='\n', strip=True)
                
                elif file_ext == '.pdf':
                    # PDF文件（需要安装PyPDF2或pdfplumber）
                    try:
                        import PyPDF2
                        with open(file_path, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            content = '\n'.join([page.extract_text() for page in pdf_reader.pages])
                    except ImportError:
                        return {'error': 'PDF解析需要安装PyPDF2: pip install PyPDF2'}, 500
                    except Exception as e:
                        return {'error': f'PDF解析失败: {str(e)}'}, 500
                
                elif file_ext in ['.docx']:
                    # Word文档（需要安装python-docx）
                    try:
                        from docx import Document as DocxDocument
                        doc = DocxDocument(file_path)
                        content = '\n'.join([para.text for para in doc.paragraphs])
                    except ImportError:
                        return {'error': 'Word文档解析需要安装python-docx: pip install python-docx'}, 500
                    except Exception as e:
                        return {'error': f'Word文档解析失败: {str(e)}'}, 500
                
                else:
                    return {'error': f'不支持的文件类型: {file_ext}'}, 400
                
                # 生成摘要（取前200字符）
                summary = content[:200] + '...' if len(content) > 200 else content
                
                # 创建文档记录
                doc = Document(
                    title=title,
                    content=content,
                    summary=summary,
                    source_type='web',  # 上传的文件标记为web类型
                    source_name='文件上传',
                    source_url=f'file://{file_path}',
                    tags=tags,
                    file_path=file_path,
                    file_size=file_size,
                    extra_metadata={
                        'original_filename': filename,
                        'file_type': file_ext,
                        'uploaded_at': datetime.utcnow().isoformat()
                    }
                )
                
                db.session.add(doc)
                db.session.commit()
                
                return {
                    'message': '文件上传成功',
                    'document': doc.to_dict_full()
                }, 201
                
            except Exception as e:
                # 如果解析失败，删除已保存的文件
                if os.path.exists(file_path):
                    os.remove(file_path)
                return {'error': f'文件解析失败: {str(e)}'}, 500
                
        except Exception as e:
            db.session.rollback()
            return {'error': f'上传失败: {str(e)}'}, 500

